#!/usr/bin/env python3
"""Convert a jetmemo markdown bench memo to a formatted .docx file.

Matches the formatting of the Court's bench memo template:
  - QTPalatine 13pt, justified, 1.2 line spacing
  - Page: 8.5x11, margins 1.25/1/1/1 inches
  - Styles: Title, Address Block, Heading 1/2/3, Main Body Text

Usage:
    python3 memo_to_docx.py input_memo.md [output.docx]

If output path is omitted, writes to {input_stem}_memo.docx.
"""

import argparse
import re
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Style constants (from example docx)
# ---------------------------------------------------------------------------
FONT_NAME = "QTPalatine"
FONT_SIZE = Pt(13)
LINE_SPACING = 1.2
NORMAL_SPACE_AFTER = Pt(6)
BODY_SPACE_AFTER = Pt(13)
BLOCK_QUOTE_INDENT = Inches(0.5)
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN_LEFT = Inches(1.25)
MARGIN_RIGHT = Inches(1)
MARGIN_TOP = Inches(1)
MARGIN_BOTTOM = Inches(1)
TAB_STOP = Inches(1)


def setup_numbering(doc):
    """Create numbering definitions matching the bench memo template.

    Returns (heading_num_id, body_num_id) for linking to paragraphs.
    """
    # Ensure numbering part exists
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # --- Abstract numbering 0: Heading multilevel list ---
    heading_abstract = parse_xml(
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="10">'
        '  <w:multiLevelType w:val="multilevel"/>'
        '  <w:lvl w:ilvl="0">'
        '    <w:start w:val="1"/>'
        '    <w:numFmt w:val="none"/>'
        '    <w:pStyle w:val="Heading1"/>'
        '    <w:suff w:val="nothing"/>'
        '    <w:lvlText w:val=""/>'
        '    <w:lvlJc w:val="center"/>'
        '    <w:pPr><w:ind w:left="0" w:firstLine="0"/></w:pPr>'
        '    <w:rPr><w:rFonts w:hint="default"/></w:rPr>'
        '  </w:lvl>'
        '  <w:lvl w:ilvl="1">'
        '    <w:start w:val="1"/>'
        '    <w:numFmt w:val="upperRoman"/>'
        '    <w:pStyle w:val="Heading2"/>'
        '    <w:lvlText w:val="%2."/>'
        '    <w:lvlJc w:val="left"/>'
        '    <w:pPr><w:ind w:left="720" w:hanging="720"/></w:pPr>'
        '    <w:rPr><w:rFonts w:hint="default"/></w:rPr>'
        '  </w:lvl>'
        '  <w:lvl w:ilvl="2">'
        '    <w:start w:val="1"/>'
        '    <w:numFmt w:val="upperLetter"/>'
        '    <w:pStyle w:val="Heading3"/>'
        '    <w:lvlText w:val="%3."/>'
        '    <w:lvlJc w:val="right"/>'
        '    <w:pPr><w:ind w:left="1440" w:hanging="720"/></w:pPr>'
        '    <w:rPr><w:rFonts w:hint="default"/></w:rPr>'
        '  </w:lvl>'
        '</w:abstractNum>'
    )
    numbering_elm.append(heading_abstract)

    # --- Abstract numbering 1: Body text [¶N] numbering ---
    body_abstract = parse_xml(
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="11">'
        '  <w:multiLevelType w:val="hybridMultilevel"/>'
        '  <w:lvl w:ilvl="0">'
        '    <w:start w:val="1"/>'
        '    <w:numFmt w:val="decimal"/>'
        '    <w:pStyle w:val="MainBodyText"/>'
        '    <w:lvlText w:val="[&#182;%1]"/>'
        '    <w:lvlJc w:val="left"/>'
        '    <w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>'
        '    <w:rPr><w:rFonts w:hint="default"/></w:rPr>'
        '  </w:lvl>'
        '</w:abstractNum>'
    )
    numbering_elm.append(body_abstract)

    # --- Concrete numbering instances ---
    heading_num = parse_xml(
        f'<w:num {nsdecls("w")} w:numId="10">'
        '  <w:abstractNumId w:val="10"/>'
        '</w:num>'
    )
    numbering_elm.append(heading_num)

    body_num = parse_xml(
        f'<w:num {nsdecls("w")} w:numId="11">'
        '  <w:abstractNumId w:val="11"/>'
        '</w:num>'
    )
    numbering_elm.append(body_num)

    return 10, 11  # heading_num_id, body_num_id


def _link_style_to_numbering(style, num_id, ilvl=0):
    """Add numPr to a style's pPr element."""
    ppr = style.element.find(qn('w:pPr'))
    if ppr is None:
        ppr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        style.element.append(ppr)
    num_pr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="{ilvl}"/>'
        f'  <w:numId w:val="{num_id}"/>'
        '</w:numPr>'
    )
    ppr.append(num_pr)


def _force_font(style, font_name):
    """Force a style's font, removing any theme font overrides."""
    style.font.name = font_name
    rpr = style.element.find(qn('w:rPr'))
    if rpr is not None:
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is not None:
            for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
                rfonts.attrib.pop(qn(attr), None)


def setup_styles(doc, heading_num_id, body_num_id):
    """Configure document styles to match the bench memo template."""
    # --- Normal ---
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    pf = normal.paragraph_format
    pf.space_after = NORMAL_SPACE_AFTER
    pf.space_before = Pt(0)
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Title ---
    title = doc.styles["Title"]
    _force_font(title, FONT_NAME)
    title.font.size = FONT_SIZE
    title.font.bold = False
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = NORMAL_SPACE_AFTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.line_spacing = LINE_SPACING

    # --- Heading 1 (centered section heads: BACKGROUND, CONCLUSION) ---
    h1 = doc.styles["Heading 1"]
    _force_font(h1, FONT_NAME)
    h1.font.size = FONT_SIZE
    h1.font.bold = False
    h1.font.color.rgb = None  # inherit (black)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_after = NORMAL_SPACE_AFTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.line_spacing = LINE_SPACING
    _link_style_to_numbering(h1, heading_num_id, ilvl=0)

    # --- Heading 2 (issue headings: I., II.) ---
    h2 = doc.styles["Heading 2"]
    _force_font(h2, FONT_NAME)
    h2.font.size = FONT_SIZE
    h2.font.bold = True
    h2.font.color.rgb = None
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h2.paragraph_format.space_after = NORMAL_SPACE_AFTER
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.line_spacing = LINE_SPACING
    _link_style_to_numbering(h2, heading_num_id, ilvl=1)

    # --- Heading 3 (sub-arguments: A., B.) ---
    h3 = doc.styles["Heading 3"]
    _force_font(h3, FONT_NAME)
    h3.font.size = FONT_SIZE
    h3.font.bold = True
    h3.font.color.rgb = None
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h3.paragraph_format.space_after = BODY_SPACE_AFTER
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.line_spacing = 1.0
    _link_style_to_numbering(h3, heading_num_id, ilvl=2)

    return doc


def add_address_block_style(doc):
    """Create the Address Block style if it doesn't exist."""
    styles = doc.styles
    try:
        ab = styles["Address Block"]
    except KeyError:
        ab = styles.add_style("Address Block", 1)  # WD_STYLE_TYPE.PARAGRAPH
        ab.base_style = styles["Normal"]
    ab.font.name = FONT_NAME
    ab.font.size = FONT_SIZE
    ab.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ab.paragraph_format.line_spacing = 1.0
    ab.paragraph_format.space_after = Pt(0)
    ab.paragraph_format.space_before = Pt(0)
    # Tab stop at 1 inch
    ab.paragraph_format.tab_stops.add_tab_stop(TAB_STOP)
    return ab


def add_body_style(doc, body_num_id):
    """Create the Main Body Text style if it doesn't exist."""
    styles = doc.styles
    try:
        mbt = styles["Main Body Text"]
    except KeyError:
        mbt = styles.add_style("Main Body Text", 1)
        mbt.base_style = styles["Normal"]
    mbt.font.name = FONT_NAME
    mbt.font.size = FONT_SIZE
    mbt.paragraph_format.space_after = BODY_SPACE_AFTER
    mbt.paragraph_format.space_before = Pt(0)
    mbt.paragraph_format.first_line_indent = Inches(0)
    mbt.paragraph_format.left_indent = Inches(0)
    mbt.paragraph_format.line_spacing = LINE_SPACING
    mbt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Link to [¶N] auto-numbering
    _link_style_to_numbering(mbt, body_num_id, ilvl=0)
    return mbt


def setup_page(doc):
    """Configure page size, margins, and footer with page numbers."""
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    # Footer with centered page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
        f'<w:r><w:t>1</w:t></w:r></w:fldSimple>'
    )
    run._element.append(parse_xml(fld_xml))


# ---------------------------------------------------------------------------
# Inline markdown parsing
# ---------------------------------------------------------------------------
# Splits text into segments: (text, bold, italic)
_INLINE_RE = re.compile(
    r"(\*\*\*(.+?)\*\*\*"   # ***bold italic***
    r"|\*\*(.+?)\*\*"       # **bold**
    r"|\*(.+?)\*"           # *italic*
    r"|([^*]+)"             # plain text
    r"|(\*+))"              # leftover asterisks
)


def add_formatted_runs(paragraph, text):
    """Parse markdown inline formatting and add runs to the paragraph."""
    for m in _INLINE_RE.finditer(text):
        if m.group(2):  # bold italic
            r = paragraph.add_run(m.group(2))
            r.bold = True
            r.italic = True
        elif m.group(3):  # bold
            r = paragraph.add_run(m.group(3))
            r.bold = True
        elif m.group(4):  # italic
            r = paragraph.add_run(m.group(4))
            r.italic = True
        elif m.group(5):  # plain
            paragraph.add_run(m.group(5))
        elif m.group(6):  # leftover asterisks
            paragraph.add_run(m.group(6))


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

def parse_header(lines):
    """Extract metadata from the markdown header block.

    Returns (metadata_dict, remaining_lines).
    """
    meta = {
        "case_number": "",
        "case_name": "",
        "argument_date": "",
        "author": "",
    }
    i = 0
    # Skip leading blank lines and the # BENCH MEMO line
    while i < len(lines) and (not lines[i].strip() or lines[i].strip().startswith("# ")):
        i += 1

    # Read bold metadata lines
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        # Bold metadata: **Key: Value** or **Value**
        bold_match = re.match(r"\*\*(.+?)\*\*", line)
        if not bold_match:
            break
        content = bold_match.group(1)
        if content.startswith("Case No."):
            meta["case_number"] = content.replace("Case No.", "").strip().rstrip(".")
        elif "Oral Argument" in content:
            meta["argument_date"] = content.split(":", 1)[1].strip() if ":" in content else ""
        elif content.startswith("Claude"):
            meta["author"] = content
        else:
            # Assume case name
            meta["case_name"] = content
        i += 1

    return meta, lines[i:]


def classify_line(line):
    """Classify a markdown line by type."""
    stripped = line.strip()
    if not stripped:
        return "blank", stripped
    if stripped.startswith("### "):
        return "h3", stripped[4:]
    if stripped.startswith("## "):
        return "h2", stripped[3:]
    if stripped.startswith("# "):
        return "h1", stripped[2:]
    if stripped.startswith("- "):
        return "bullet", stripped[2:]
    if stripped.startswith("|"):
        return "table_row", stripped
    if stripped.startswith("> "):
        return "blockquote", stripped[2:]
    return "body", stripped


def is_section_heading(text):
    """Check if an h2 is a major section (BACKGROUND, CONCLUSION) vs issue heading."""
    upper = text.strip().upper()
    return upper in ("BACKGROUND", "CONCLUSION", "DISCUSSION")


def build_table(doc, rows_text):
    """Build a docx table from markdown table rows."""
    # Parse rows
    parsed = []
    for row_text in rows_text:
        cells = [c.strip() for c in row_text.strip("|").split("|")]
        parsed.append(cells)

    if len(parsed) < 2:
        return

    # Skip separator row (contains dashes)
    header = parsed[0]
    data_rows = [r for r in parsed[1:] if not all(re.match(r"^[-:]+$", c) for c in r)]

    if not data_rows:
        return

    ncols = len(header)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.style = doc.styles["Normal"]
        r = p.add_run(cell_text)
        r.bold = True

    # Data rows
    for i, row_data in enumerate(data_rows):
        for j in range(min(ncols, len(row_data))):
            cell = table.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            p.style = doc.styles["Normal"]
            add_formatted_runs(p, row_data[j])

    # Style table with light borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f"<w:tblPr {nsdecls('w')}/>")
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders)


def strip_para_number(text):
    """Remove [¶N] prefix from paragraph text, return (text, para_num)."""
    m = re.match(r"\[¶(\d+)\]\s*", text)
    if m:
        return text[m.end():], m.group(1)
    return text, None


def strip_heading_number(text):
    """Remove manual numbering prefix from heading text.

    Strips Roman numeral prefixes (I., II., III.) from H2 headings and
    letter prefixes (A., B., C.) from H3 headings, since Word's numbering
    definitions supply these automatically.
    """
    # Roman numeral prefix: "I. ", "II. ", "IV. ", etc.
    m = re.match(r"^[IVXLC]+\.\s+", text)
    if m:
        return text[m.end():]
    # Letter prefix: "A. ", "B. ", etc.
    m = re.match(r"^[A-Z]\.\s+", text)
    if m:
        return text[m.end():]
    return text


def _suppress_numbering(paragraph):
    """Override paragraph numbering with numId=0 to suppress auto-number."""
    ppr = paragraph._element.find(qn('w:pPr'))
    if ppr is None:
        ppr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        paragraph._element.insert(0, ppr)
    # Remove existing numPr if present
    existing = ppr.find(qn('w:numPr'))
    if existing is not None:
        ppr.remove(existing)
    num_pr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        '  <w:ilvl w:val="0"/>'
        '  <w:numId w:val="0"/>'
        '</w:numPr>'
    )
    ppr.append(num_pr)


def convert(md_path, docx_path):
    """Convert a markdown bench memo to docx."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    lines = md_text.split("\n")

    # Parse header metadata
    meta, body_lines = parse_header(lines)

    # Create document and set up styles
    doc = Document()
    heading_num_id, body_num_id = setup_numbering(doc)
    setup_styles(doc, heading_num_id, body_num_id)
    add_address_block_style(doc)
    add_body_style(doc, body_num_id)
    setup_page(doc)

    # Remove the default empty paragraph
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # --- Title ---
    p = doc.add_paragraph("MEMORANDUM", style="Title")

    # --- Address Block ---
    today_str = date.today().strftime("%B %d, %Y").replace(" 0", " ")

    addr_lines = [
        ("To:", "Justice Tufte"),
        ("From:", meta["author"] or "Claude (AI first draft)"),
        ("Date:", today_str),
    ]
    # Re line
    re_text = meta["case_name"]
    if meta["case_number"]:
        re_text += f" - No. {meta['case_number']}"
    addr_lines.append(("Re:", re_text))
    if meta["argument_date"]:
        addr_lines.append(("Argument:", meta["argument_date"]))

    for label, value in addr_lines:
        p = doc.add_paragraph(style="Address Block")
        p.add_run(label)
        p.add_run("\t")
        # Italicize case name in Re: line
        if label == "Re:" and meta["case_name"]:
            r = p.add_run(meta["case_name"])
            r.italic = True
            if meta["case_number"]:
                p.add_run(f" - No. {meta['case_number']}")
        else:
            p.add_run(value)

    # --- Body ---
    i = 0
    in_quick_ref = False
    table_rows = []

    while i < len(body_lines):
        line_type, content = classify_line(body_lines[i])

        # Handle table accumulation
        if line_type == "table_row":
            table_rows.append(content)
            i += 1
            continue
        elif table_rows:
            build_table(doc, table_rows)
            table_rows = []

        if line_type == "blank":
            i += 1
            continue

        if line_type == "h1":
            # Top-level heading (shouldn't appear in body normally)
            p = doc.add_paragraph(content, style="Heading 1")
            i += 1
            continue

        if line_type == "h2":
            in_quick_ref = content.strip().lower() == "quick reference"
            if content.strip().lower().startswith("key exhibits"):
                # Key Exhibits heading - use Heading 2
                p = doc.add_paragraph(strip_heading_number(content), style="Heading 2")
                i += 1
                continue
            if in_quick_ref:
                # Skip the Quick Reference heading in docx; items become bullets
                i += 1
                continue
            if is_section_heading(content):
                p = doc.add_paragraph(content, style="Heading 1")
            else:
                p = doc.add_paragraph(strip_heading_number(content), style="Heading 2")
            i += 1
            continue

        if line_type == "h3":
            p = doc.add_paragraph(strip_heading_number(content), style="Heading 3")
            in_quick_ref = False
            i += 1
            continue

        if line_type == "bullet":
            # Accumulate multi-line bullets
            text = content
            i += 1
            while i < len(body_lines):
                next_stripped = body_lines[i].strip()
                if next_stripped and not next_stripped.startswith("- ") and not next_stripped.startswith("#") and not next_stripped.startswith("|"):
                    # Continuation line
                    if next_stripped.startswith("[¶") or not body_lines[i].startswith("  "):
                        break
                    text += " " + next_stripped
                    i += 1
                else:
                    break
            # In Quick Reference, use body text with a bullet character
            p = doc.add_paragraph(style="Main Body Text")
            _suppress_numbering(p)
            p.paragraph_format.space_after = NORMAL_SPACE_AFTER
            add_formatted_runs(p, "• " + text)
            continue

        if line_type == "blockquote":
            text, _pn = strip_para_number(content)
            p = doc.add_paragraph(style="Main Body Text")
            _suppress_numbering(p)
            p.paragraph_format.left_indent = BLOCK_QUOTE_INDENT
            p.paragraph_format.first_line_indent = BLOCK_QUOTE_INDENT
            p.paragraph_format.right_indent = BLOCK_QUOTE_INDENT
            p.paragraph_format.line_spacing = 1.0
            add_formatted_runs(p, text)
            i += 1
            continue

        if line_type == "body":
            in_quick_ref = False
            text, _pn = strip_para_number(content)
            # Accumulate continuation lines (non-blank, non-heading, non-bullet)
            i += 1
            while i < len(body_lines):
                next_type, next_content = classify_line(body_lines[i])
                if next_type == "blank" or next_type.startswith("h") or next_type == "bullet" or next_type == "table_row":
                    break
                # Check if next line starts a new paragraph marker
                if re.match(r"\[¶\d+\]", next_content):
                    break
                text += " " + next_content
                i += 1
            p = doc.add_paragraph(style="Main Body Text")
            add_formatted_runs(p, text)
            continue

        i += 1

    # Flush any remaining table rows
    if table_rows:
        build_table(doc, table_rows)

    # Save
    doc.save(str(docx_path))
    return docx_path


def main():
    parser = argparse.ArgumentParser(description="Convert jetmemo markdown to docx")
    parser.add_argument("input", help="Path to markdown memo file")
    parser.add_argument("output", nargs="?", help="Output .docx path (default: {input_stem}.docx)")
    args = parser.parse_args()

    md_path = Path(args.input)
    if not md_path.exists():
        print(f"Error: {md_path} not found", file=sys.stderr)
        sys.exit(1)

    if args.output:
        docx_path = Path(args.output)
    else:
        docx_path = md_path.with_suffix(".docx")

    out = convert(md_path, docx_path)
    print(f"Written: {out}")


if __name__ == "__main__":
    main()
